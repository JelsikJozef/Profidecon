from .base import BaseParser, ParsedDocument
from docx import Document as DocxDocument
from pathlib import Path

class DocxParser(BaseParser):
    suffixes = (".docx",)

    def parse(self, path: Path) -> ParsedDocument:
        doc = DocxDocument(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        return ParsedDocument(text=text, metadata={"source": str(path)})
